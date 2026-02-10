import os
import time
from datetime import datetime
from dotenv import load_dotenv
from github import Github
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
import chromadb
from chromadb.config import Settings
from openai import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
import pdfplumber
from docx import Document

# ================= CONFIG =================
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
GITHUB_USERNAME = os.getenv("GITHUB_USERNAME") 

DATA_DIR = Path("data")
CHROMA_DIR = Path("chroma_db")
TEMP_DIR = Path("temp_github")

if not OPENAI_API_KEY or not GITHUB_TOKEN:
    raise ValueError("❌ Missing OPENAI_API_KEY or GITHUB_TOKEN in .env")

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(CHROMA_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)

# Initialize OpenAI client 
openai_client = OpenAI(api_key=OPENAI_API_KEY)

# Initialize ChromaDB client
chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))

# Text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    length_function=len,
)

def get_embeddings(texts):
    """Get embeddings using OpenAI directly"""
    response = openai_client.embeddings.create(
        model="text-embedding-3-small",
        input=texts
    )
    return [item.embedding for item in response.data]

# ================= FILE PROCESSING =================
def extract_text(file_path):
    """Extract text based on file extension"""
    text = ""
    suffix = file_path.suffix.lower()
    try:
        if suffix in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif suffix == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
        elif suffix == ".docx":
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        elif suffix == ".doc":
            import textract
            text = textract.process(str(file_path)).decode('utf-8')
    except Exception as e:
        print(f"⚠️ Error extracting {file_path.name}: {e}")
    return text

def process_single_file(file_path):
    """Process and embed a single file if updated"""
    print(f"\n📄 Checking file: {file_path.name}")
    
    # Get collection
    try:
        collection = chroma_client.get_collection("cv_collection")
    except:
        collection = chroma_client.create_collection(name="cv_collection")
        
    # Check if needs update
    last_modified = os.path.getmtime(file_path)
    existing = collection.get(where={"filename": file_path.name}, include=['metadatas'])
    
    if existing['metadatas']:
        # Check stored timestamp
        stored_modified = existing['metadatas'][0].get('last_modified', 0)
        if last_modified <= stored_modified:
            print(f"  ⏭️  File up to date: {file_path.name}")
            return
            
        print(f"  🔄 Updating file: {file_path.name}")
        # Delete old chunks
        collection.delete(where={"filename": file_path.name})
    else:
        print(f"  🆕 New file: {file_path.name}")

    # Extract and embed
    content = extract_text(file_path)
    if not content.strip():
        print(f"  ⚠️ No content extracted from {file_path.name}")
        return

    chunks = text_splitter.split_text(content)
    embeddings = get_embeddings(chunks)
    
    ids = [f"{file_path.name}_{i}" for i in range(len(chunks))]
    metadatas = [{
        "filename": file_path.name,
        "chunk_index": i,
        "last_modified": last_modified,
        "source": "local_file"
    } for i in range(len(chunks))]
    
    collection.add(
        embeddings=embeddings,
        documents=chunks,
        metadatas=metadatas,
        ids=ids
    )
    print(f"  ✅ Embedded {len(chunks)} chunks from {file_path.name}")

# ================= GITHUB PROCESSING =================
def process_single_repo(repo_name):
    """Process and embed a single repo if updated"""
    print(f"\n📦 Checking repo: {repo_name}")
    g = Github(GITHUB_TOKEN)
    
    try:
        user = g.get_user(GITHUB_USERNAME) if GITHUB_USERNAME else g.get_user()
        repo = user.get_repo(repo_name)
        
        # Get collection
        try:
            collection = chroma_client.get_collection("github_collection")
        except:
            collection = chroma_client.create_collection(name="github_collection")
            
        # Check if needs update
        pushed_at = repo.pushed_at.timestamp()
        existing = collection.get(where={"repo_name": repo_name}, include=['metadatas'])
        
        if existing['metadatas']:
            stored_pushed = existing['metadatas'][0].get('pushed_at', 0)
            if pushed_at <= stored_pushed:
                print(f"  ⏭️  Repo up to date: {repo_name}")
                return
            print(f"  🔄 Updating repo: {repo_name}")
            collection.delete(where={"repo_name": repo_name})
        else:
            print(f"  🆕 New repo: {repo_name}")

        # Fetch Content
        content_parts = []
        
        # 1. README
        try:
            readme = repo.get_readme().decoded_content.decode("utf-8")
            content_parts.append(f"README:\n{readme}")
        except:
            content_parts.append("README: Not available")
            
        # 2. Tech Stack (Package files)
        tech_stack = []
        try:
            contents = repo.get_contents("")
            file_names = [c.name for c in contents]
            if "requirements.txt" in file_names: tech_stack.append("Python")
            if "package.json" in file_names: tech_stack.append("Node.js/JavaScript")
            if "pom.xml" in file_names: tech_stack.append("Java")
            if "go.mod" in file_names: tech_stack.append("Go")
            if "Cargo.toml" in file_names: tech_stack.append("Rust")
            if "Dockerfile" in file_names: tech_stack.append("Docker")
        except:
            pass
            
        full_content = f"""
Repository: {repo.name}
URL: {repo.html_url}
Description: {repo.description or "No description"}
Language: {repo.language}
Stars: {repo.stargazers_count}
Tech Stack Detected: {', '.join(tech_stack)}

{chr(10).join(content_parts)}
"""
        chunks = text_splitter.split_text(full_content)
        embeddings = get_embeddings(chunks)
        
        ids = [f"{repo_name}_{i}" for i in range(len(chunks))]
        metadatas = [{
            "repo_name": repo_name,
            "chunk_index": i,
            "pushed_at": pushed_at,
            "source": "github"
        } for i in range(len(chunks))]
        
        collection.add(
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )
        print(f"  ✅ Embedded {len(chunks)} chunks from {repo_name}")
        
    except Exception as e:
        print(f"❌ Error processing repo {repo_name}: {e}")

# ================= MAIN SYNC =================
def sync_all():
    """Sync all files and repos"""
    print("🚀 Starting sync...")
    
    # 1. Sync Files
    print("\n📂 Syncing local files...")
    if DATA_DIR.exists():
        for file_path in DATA_DIR.glob("*"):
            if file_path.suffix.lower() in [".txt", ".md", ".pdf", ".docx", ".doc"]:
                process_single_file(file_path)
    
    # 2. Sync Repos
    print("\n🔍 Syncing GitHub repos...")
    g = Github(GITHUB_TOKEN)
    try:
        user = g.get_user(GITHUB_USERNAME) if GITHUB_USERNAME else g.get_user()
        for repo in user.get_repos():
            process_single_repo(repo.name)
    except Exception as e:
        print(f"❌ Error fetching repos: {e}")

if __name__ == "__main__":
    sync_all()